"""
docx_to_oscal.py — Workshop Edition (Discovery-Driven)

Converts a FedRAMP/CUI-style Word SSP into OSCAL JSON.
Reads control tables from a .docx file and produces the same
separated OSCAL artifacts as excel_to_oscal.py.

Components are extracted from SSP narratives via keyword matching.
Discovery fills in the real environment inventory.

SUPPORTED SSP PATTERNS:
  - CUI template: requirement text -> status checkboxes -> narrative
  - FedRAMP hybrid: navy header bar -> requirement -> status row -> narrative
  - Any Word SSP with per-control table blocks (auto-detected)

ARCHITECTURE:
  Same as excel_to_oscal.py:
  - OSCAL skeleton is master. SSP carries the CLAIM.
  - Evidence goes into assessment-results.json (separate artifact).
  - Ingest scripts write there -- not here.

USAGE:
  python docx_to_oscal.py --input ssp.docx --output oscal/
  python docx_to_oscal.py --help

MODELS PRODUCED:
  Model 5: oscal/ssp.json              -- System Security Plan (the claim)
  Model 7: oscal/assessment-results.json -- skeleton (ingest scripts populate)
  Model 8: oscal/poam.json              -- skeleton (reconciler populates)
"""

import json
import os
import sys
import re
import uuid
import argparse
from datetime import datetime, timezone
from pathlib import Path

try:
    from docx import Document
except ImportError:
    print("ERROR: python-docx not installed.")
    print("Run: pip install python-docx")
    sys.exit(1)

# Import shared utilities
sys.path.insert(0, os.path.dirname(__file__))
from pipeline_utils import (
    stable_uuid, OSCAL_NAMESPACE, KNOWN_COMPONENTS, extract_components_from_text,
)


# ── Status detection ──────────────────────────────────────────────────────────
# The CUI/FedRAMP pattern uses checkbox characters in table cells.
# ☑ (U+2611) = checked, ☐ (U+2610) = unchecked

VALID_STATUSES = {
    "implemented":    "implemented",
    "inherited":      "inherited",
    "planned":        "planned",
    "not applicable": "not-applicable",
    "not-applicable": "not-applicable",
    "n/a":            "not-applicable",
}


# ── Helpers ───────────────────────────────────────────────────────────────────

def normalize_control_id(raw_id: str) -> str:
    """Normalize control IDs: AC-01 -> ac-1, AC-2(01) -> ac-2(1)"""
    if not raw_id:
        return None
    s = str(raw_id).strip().lower()
    s = re.sub(r'-0*(\d)', r'-\1', s)
    s = re.sub(r'\(0*(\d+)\)', r'(\1)', s)
    return s


def extract_control_id_and_name(header_text: str):
    """
    Parse the header row: 'AC-2  Account Management' or
    'AC-2(1)  Account Management | Automated System Account Management'
    Returns (control_id, control_name) or (None, None).
    """
    header_text = header_text.strip()
    # Match patterns like AC-1, AC-2(1), PE-3, SI-5(1)
    m = re.match(r'^([A-Za-z]{2}-\d+(?:\(\d+\))?)\s+(.+)$', header_text)
    if m:
        return m.group(1).strip(), m.group(2).strip()
    return None, None


def detect_status_from_checkboxes(cells):
    """
    Read the CUI-style checkbox row. Each cell contains ☑ or ☐ followed
    by the status label. Returns the checked status.
    """
    for cell_text in cells:
        text = cell_text.strip()
        if '\u2611' in text:  # ☑ = checked
            # Extract the label after the checkbox
            label = text.replace('\u2611', '').replace('\u2610', '').strip()
            key = label.lower()
            return VALID_STATUSES.get(key, "implemented")
    return "not-implemented"


def extract_requirement_text(cell_text: str) -> str:
    """Strip the 'Requirement:' prefix from the requirement cell."""
    text = cell_text.strip()
    if text.lower().startswith("requirement:"):
        text = text[len("requirement:"):].strip()
    return text


def is_missing_or_stale(text, status):
    if not text:
        return True, "missing"
    if status == "implemented" and len(text) < 50:
        return True, "stale"
    return False, None


# ── OSCAL builders (discovery-driven) ────────────────────────────────────────

def build_by_components_from_narrative(narrative: str) -> tuple:
    """
    Scan the narrative for mentions of known tools/services.
    Returns (by_components list, component_keys found).
    No control mapping — just name recognition.
    """
    component_keys = extract_components_from_text(narrative)
    by_components = []
    for key in component_keys:
        comp = KNOWN_COMPONENTS[key]
        by_components.append({
            "component-uuid": stable_uuid(f"component:{key}"),
            "description": comp["title"],
            "implementation-status": {
                "state": "claimed",
                "remarks": "referenced in SSP narrative"
            },
            "props": [
                {"name": "component-key", "value": key},
                {"name": "origin",        "value": "ssp-narrative"},
            ]
        })
    return by_components, component_keys


def build_component_definitions(discovered_keys: set):
    """Build components section from what was actually found in narratives."""
    components = [
        {
            "uuid": stable_uuid("component:this-system"),
            "type": "this-system",
            "title": "Workshop Demo System",
            "description": (
                "A demo system built during the GRC Engineering Club "
                "OSCAL builder session."
            ),
            "status": {"state": "operational"}
        }
    ]
    for key in sorted(discovered_keys):
        comp = KNOWN_COMPONENTS[key]
        components.append({
            "uuid": stable_uuid(f"component:{key}"),
            "type": comp["type"],
            "title": comp["title"],
            "props": [
                {"name": "component-key", "value": key},
                {"name": "origin",        "value": "ssp-narrative"},
            ],
            "status": {"state": "operational"}
        })
    return components


def build_assessment_results_skeleton(ssp_uuid: str):
    ar_uuid = stable_uuid("assessment-results:workshop")
    return {
        "assessment-results": {
            "uuid": ar_uuid,
            "metadata": {
                "title": "Workshop Demo \u2014 Assessment Results",
                "last-modified": datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ"),
                "version": "1.0.0",
                "oscal-version": "1.1.2",
            },
            "import-ap": {
                "href": "#",
                "remarks": "Assessment plan is automated \u2014 the pipeline scripts define what's checked and how."
            },
            "results": [
                {
                    "uuid": stable_uuid("result:workshop-run"),
                    "title": "Workshop Pipeline Run",
                    "description": "Evidence collected against FedRAMP Moderate baseline.",
                    "start": datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ"),
                    "reviewed-controls": {
                        "control-selections": [
                            {"include-all": {}}
                        ]
                    },
                    "observations": [],
                    "findings": [],
                }
            ]
        }
    }


def build_poam_skeleton(ssp_uuid: str):
    return {
        "plan-of-action-and-milestones": {
            "uuid": stable_uuid("poam:workshop"),
            "metadata": {
                "title": "Workshop Demo \u2014 Plan of Action and Milestones",
                "last-modified": datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ"),
                "version": "1.0.0",
                "oscal-version": "1.1.2",
            },
            "import-ssp": {"href": "ssp.json"},
            "poam-items": []
        }
    }


# ── Word SSP Parser ──────────────────────────────────────────────────────────

def parse_control_table(table):
    """
    Parse a single control table block from the Word SSP.

    Expected structure (CUI/FedRAMP hybrid pattern):
      Row 0: Control ID + Name (header bar)
      Row 1: Requirement text
      Row 2: Status checkboxes (Implemented | Planned | N/A | Inherited)
      Row 3: Evidence Method + Origination
      Row 4: Implementation Description header
      Row 5: Implementation narrative

    Returns dict or None if not a control table.
    """
    rows = table.rows
    if len(rows) < 4:
        return None

    # Row 0: Header — extract control ID and name
    header_text = rows[0].cells[0].text.strip()
    control_id_raw, control_name = extract_control_id_and_name(header_text)
    if not control_id_raw:
        return None

    control_id = normalize_control_id(control_id_raw)

    # Row 1: Requirement text
    requirement = ""
    if len(rows) > 1:
        requirement = extract_requirement_text(rows[1].cells[0].text)

    # Row 2: Status checkboxes
    status = "not-implemented"
    if len(rows) > 2:
        cell_texts = [c.text for c in rows[2].cells]
        # De-duplicate (merged cells repeat text)
        seen = set()
        unique = []
        for ct in cell_texts:
            if ct not in seen:
                seen.add(ct)
                unique.append(ct)
        status = detect_status_from_checkboxes(unique)

    # Row 5 (or last row): Implementation narrative
    narrative = ""
    if len(rows) > 5:
        narrative = rows[5].cells[0].text.strip()
    elif len(rows) > 3:
        # Fallback: last row is the narrative
        narrative = rows[-1].cells[0].text.strip()
        # Skip if it's just the header text
        if narrative.lower() in ("implementation description", ""):
            narrative = ""

    return {
        "control_id": control_id,
        "control_name": control_name,
        "requirement": requirement,
        "status": status,
        "narrative": narrative,
    }


# ── Main converter ────────────────────────────────────────────────────────────

def convert_docx_to_oscal(input_path: str, output_dir: str):
    print(f"\n{'='*62}")
    print(f"  Word SSP \u2192 OSCAL Converter")
    print(f"  Discovery-Driven \u2014 SSP as claim, no assumptions")
    print(f"  Reads FedRAMP/CUI-style .docx control tables")
    print(f"{'='*62}")
    print(f"  Input:      {input_path}")
    print(f"  Output dir: {output_dir}")
    print(f"{'='*62}\n")

    print(f"  Opening Word SSP...")
    try:
        doc = Document(input_path)
    except Exception as e:
        print(f"  ERROR: Could not open Word file: {e}")
        return False

    print(f"  Found {len(doc.tables)} tables in document")
    print(f"  Scanning for control tables and extracting components...\n")

    # Parse all control tables
    controls = []
    for table in doc.tables:
        ctrl = parse_control_table(table)
        if ctrl:
            controls.append(ctrl)

    if not controls:
        print("  ERROR: No control tables found in document.")
        print("  Expected format: tables with control ID (e.g., AC-1) in first row.")
        return False

    print(f"  Found {len(controls)} controls\n")

    # Build OSCAL
    implemented_requirements = []
    all_component_keys = set()
    stats = {
        "total": 0, "implemented": 0, "inherited": 0, "planned": 0,
        "not_applicable": 0, "has_narrative": 0, "missing_narrative": 0,
        "controls_with_components": 0, "controls_without_components": 0,
    }

    for ctrl in controls:
        stats["total"] += 1
        control_id = ctrl["control_id"]
        status = ctrl["status"]
        ssp_text = ctrl["narrative"]
        needs_review, review_reason = is_missing_or_stale(ssp_text, status)

        # Stats
        stat_map = {"implemented": "implemented", "inherited": "inherited",
                    "planned": "planned", "not-applicable": "not_applicable"}
        if status in stat_map:
            stats[stat_map[status]] += 1
        if ssp_text and not needs_review:
            stats["has_narrative"] += 1
        else:
            stats["missing_narrative"] += 1

        # Extract components from narrative
        by_components, component_keys = build_by_components_from_narrative(ssp_text or "")
        all_component_keys.update(component_keys)

        if component_keys:
            stats["controls_with_components"] += 1
        else:
            stats["controls_without_components"] += 1

        family = control_id.split("-")[0] if "-" in control_id else ""

        props = [
            {"name": "control-origination", "value": status},
            {"name": "control-family",      "value": family.upper()},
            {"name": "last-reconciled",     "value": "never"},
        ]
        if needs_review:
            props.append({"name": "review-flag", "value": review_reason})
        if ssp_text:
            props.append({"name": "baseline-narrative", "value": ssp_text})

        impl_req = {
            "uuid": stable_uuid(f"impl-req:{control_id}"),
            "control-id": control_id,
            "props": props,
            "statements": [
                {
                    "statement-id": f"{control_id}_smt",
                    "uuid": stable_uuid(f"stmt:{control_id}"),
                    "description": ssp_text or "Implementation statement not yet documented.",
                    "by-components": by_components,
                }
            ]
        }
        implemented_requirements.append(impl_req)

        comp_display = ", ".join(component_keys) if component_keys else "(none)"
        print(f"  {control_id.upper():10s} {status:18s} components: {comp_display:30s} "
              f"{'OK' if ssp_text else 'XX'} narrative")

    # ── Build the SSP (Model 5) ──────────────────────────────────────────────

    ssp_uuid = stable_uuid("ssp:workshop-demo")

    oscal_ssp = {
        "system-security-plan": {
            "uuid": ssp_uuid,
            "metadata": {
                "title": "Workshop Demo \u2014 System Security Plan",
                "last-modified": datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ"),
                "version": "1.0.0",
                "oscal-version": "1.1.2",
                "remarks": (
                    "Generated by GRC Engineering Club workshop converter (docx_to_oscal.py). "
                    "FedRAMP Moderate baseline. Components extracted from SSP "
                    "narratives \u2014 discovery fills in the real inventory."
                ),
            },
            "import-profile": {
                "href": "https://raw.githubusercontent.com/usnistgov/oscal-content/main/nist.gov/SP800-53/rev5/json/NIST_SP-800-53_rev5_MODERATE-baseline_profile.json",
                "remarks": "FedRAMP Moderate baseline from NIST OSCAL content repository"
            },
            "system-characteristics": {
                "system-ids": [{"id": "WORKSHOP-DEMO-001"}],
                "system-name": "GRC Engineering Club Workshop Demo",
                "description": (
                    "A demonstration system built during the live builder session. "
                    "Components are extracted from SSP narratives. Discovery identifies "
                    "the actual environment inventory."
                ),
                "security-sensitivity-level": "moderate",
                "system-information": {
                    "information-types": [
                        {
                            "uuid": stable_uuid("info-type:demo"),
                            "title": "Workshop Demonstration Data",
                            "description": "Non-sensitive demonstration data for OSCAL pipeline training.",
                            "categorizations": [
                                {"system": "https://doi.org/10.6028/NIST.SP.800-60v2r1"}
                            ],
                            "confidentiality-impact": {"base": "fips-199-moderate"},
                            "integrity-impact":      {"base": "fips-199-moderate"},
                            "availability-impact":   {"base": "fips-199-moderate"},
                        }
                    ]
                },
                "security-impact-level": {
                    "security-objective-confidentiality": "fips-199-moderate",
                    "security-objective-integrity":       "fips-199-moderate",
                    "security-objective-availability":    "fips-199-moderate",
                },
                "status": {"state": "operational"},
                "authorization-boundary": {
                    "description": "AWS account, GitHub repositories, and associated security tooling."
                },
            },
            "system-implementation": {
                "users": [
                    {
                        "uuid": stable_uuid("user:workshop-participant"),
                        "title": "Workshop Participant",
                        "role-ids": ["system-owner"],
                        "description": "GRC Engineering Club member building the demo pipeline."
                    }
                ],
                "components": build_component_definitions(all_component_keys),
            },
            "control-implementation": {
                "description": (
                    "FedRAMP Moderate baseline controls with implementation narratives "
                    "as documented claims. Components extracted from narratives via "
                    "keyword matching. Discovery adds the real environment inventory."
                ),
                "implemented-requirements": implemented_requirements,
            }
        }
    }

    # ── Write all three OSCAL files ──────────────────────────────────────────

    output_path = Path(output_dir)
    output_path.mkdir(parents=True, exist_ok=True)

    ssp_file = output_path / "ssp.json"
    with open(ssp_file, "w", encoding="utf-8") as f:
        json.dump(oscal_ssp, f, indent=2, ensure_ascii=False)

    ar = build_assessment_results_skeleton(ssp_uuid)
    ar_file = output_path / "assessment-results.json"
    with open(ar_file, "w", encoding="utf-8") as f:
        json.dump(ar, f, indent=2, ensure_ascii=False)

    poam = build_poam_skeleton(ssp_uuid)
    poam_file = output_path / "poam.json"
    with open(poam_file, "w", encoding="utf-8") as f:
        json.dump(poam, f, indent=2, ensure_ascii=False)

    # ── Summary ──────────────────────────────────────────────────────────────

    print(f"\n{'='*62}")
    print(f"  CONVERSION COMPLETE")
    print(f"{'='*62}")
    print(f"  Source format:            Word SSP (.docx)")
    print(f"  OSCAL SSP:                {ssp_file}")
    print(f"  Assessment Results:       {ar_file}")
    print(f"  POA&M:                    {poam_file}")
    print(f"{'─'*62}")
    print(f"  Total controls:           {stats['total']}")
    print(f"  Implemented:              {stats['implemented']}")
    print(f"  Inherited:                {stats.get('inherited', 0)}")
    print(f"  Narratives present:       {stats['has_narrative']}")
    print(f"  Missing narratives:       {stats['missing_narrative']}")
    print(f"{'─'*62}")
    print(f"  Components from SSP:      {len(all_component_keys)}")
    if all_component_keys:
        print(f"    {', '.join(KNOWN_COMPONENTS[k]['title'] for k in sorted(all_component_keys))}")
    print(f"  Controls with components: {stats['controls_with_components']}")
    print(f"  Controls without:         {stats['controls_without_components']}")
    print(f"{'─'*62}")
    print(f"  UUID strategy:            v5 deterministic (stable diffs)")
    print(f"  Architecture:             SSP = claim | Discovery = truth | AR = evidence")
    print(f"{'='*62}")
    print(f"\n  OSCAL output is identical whether source is Excel or Word.")
    print(f"  Same UUIDs, same structure, same downstream pipeline.")
    print(f"{'='*62}\n")

    return True


# ── CLI ───────────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    parser = argparse.ArgumentParser(
        description="Convert FedRAMP/CUI-style Word SSP (.docx) to OSCAL JSON"
    )
    parser.add_argument(
        "--input", "-i",
        required=True,
        help="Path to Word SSP file (.docx)"
    )
    parser.add_argument(
        "--output", "-o",
        default="oscal",
        help="Output directory for OSCAL JSON files (default: oscal/)"
    )
    args = parser.parse_args()

    success = convert_docx_to_oscal(args.input, args.output)
    sys.exit(0 if success else 1)
